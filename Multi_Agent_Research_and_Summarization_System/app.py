import os
import streamlit as st
import pdfplumber
from docx import Document as DocxDocument

from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph
from langchain_core.runnables import RunnableLambda

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchRun
from dotenv import load_dotenv
load_dotenv()
# ---------------------- CONFIGURATION ----------------------
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
search = DuckDuckGoSearchRun()

# ---------------------- FILE PARSER ----------------------
def extract_text_from_local_path(path):
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

# ---------------------- AGENTS ----------------------
def router_agent(state):
    query = state.get("query", "")
    route_prompt = PromptTemplate.from_template(
        "Classify the query into one of [web, rag, llm]:\n\nQuery: {query}\n\nAnswer:"
    )
    route_result = (route_prompt | llm).invoke({"query": query}).content.lower()
    route = "llm"
    if "web" in route_result:
        route = "web"
    elif "rag" in route_result:
        route = "rag"
    return {**state, "route": route}

def web_agent(state):
    query = state["query"]
    try:
        result = search.run(query)
        return {**state, "content": result}
    except Exception as e:
        return {**state, "content": f"Web search failed: {str(e)}"}

def rag_agent(state):
    query = state["query"]
    retriever = state["retriever"]
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    answer = qa_chain.run(query)
    return {**state, "content": answer}

def llm_agent(state):
    query = state["query"]
    response = llm.invoke(query)
    return {**state, "content": response.content}

def summarizer_agent(state):
    content = state["content"]
    prompt = PromptTemplate.from_template("Summarize clearly and concisely:\n\n{content}")
    summary = (prompt | llm).invoke({"content": content}).content
    return {**state, "final": summary}

# ---------------------- LANGGRAPH ----------------------
def run_langgraph(user_query, retriever):
    workflow = StateGraph(dict)
    workflow.set_entry_point("router")

    workflow.add_node("router", RunnableLambda(router_agent))
    workflow.add_node("web", RunnableLambda(web_agent))
    workflow.add_node("rag", RunnableLambda(rag_agent))
    workflow.add_node("llm", RunnableLambda(llm_agent))
    workflow.add_node("summarizer", RunnableLambda(summarizer_agent))

    def router_logic(state): return state["route"]
    workflow.add_conditional_edges("router", router_logic, {
        "web": "web",
        "rag": "rag",
        "llm": "llm"
    })

    for node in ["web", "rag", "llm"]:
        workflow.add_edge(node, "summarizer")

    workflow.set_finish_point("summarizer")
    app = workflow.compile()
    return app.invoke({"query": user_query, "retriever": retriever})["final"]

# ---------------------- STREAMLIT APP ----------------------
st.set_page_config(page_title="üîç Fully Agentic Research Assistant", layout="centered")
st.title("üß† Multi-Agent RAG System (LangGraph + Web + RAG + LLM)")

retriever = None
documents_loaded = False

# Load local documents
if os.path.exists("my_docs"):
    with st.spinner("üìÇ Loading documents from 'my_docs' folder..."):
        all_content = []
        for filename in os.listdir("my_docs"):
            filepath = os.path.join("my_docs", filename)
            if filename.lower().endswith(('.pdf', '.txt', '.docx')):
                content = extract_text_from_local_path(filepath)
                if content:
                    all_content.append(content)

        if all_content:
            chunks = text_splitter.create_documents(all_content)
            vectorstore = FAISS.from_documents(chunks, embeddings)
            retriever = vectorstore.as_retriever()
            documents_loaded = True
            st.success(f"‚úÖ Loaded {len(all_content)} documents.")
        else:
            st.warning("‚ö†Ô∏è No readable files found.")

if not documents_loaded:
    st.info("üìÑ Using fallback knowledge base.")
    docs = [
        Document(page_content="LangGraph is a Python framework for agent workflows."),
        Document(page_content="Gemini 1.5 Flash is fast and great for summarization."),
    ]
    vectorstore = FAISS.from_documents(docs, embeddings)
    retriever = vectorstore.as_retriever()

# User Input
query = st.text_input("üí¨ Ask your question", placeholder="e.g. What is LangGraph?")
submit = st.button("Submit")

if submit:
    if not query.strip():
        st.warning("‚ö†Ô∏è Please enter a question.")
    else:
        with st.spinner("ü§ñ Thinking..."):
            try:
                answer = run_langgraph(query, retriever)
                st.success("‚úÖ Done!")
                st.subheader("üìò Answer:")
                st.write(answer)
            except Exception as e:
                st.error(f"‚ùå Error: {str(e)}")
